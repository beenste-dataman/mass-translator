import os
import argparse
import sys
from tqdm import tqdm
import xlrd
import xlwt
# Required librarie
import docx
import openpyxl
import csv
import json
import torch
import chardet
import os
import subprocess
import email
import magic
import re
from email import policy
from bs4 import BeautifulSoup
from transformers import MarianMTModel, MarianTokenizer
from concurrent.futures import ThreadPoolExecutor
from odf.opendocument import OpenDocumentText
from odf.text import P, Span
from odf import teletype
from striprtf import rtf_to_text
from docx import Document
from striprtf import *







art = '''
___________               
\_   _____/__________     
 |    __)/  _ \_  __ \    
 |     \(  <_> )  | \/    
 \___  / \____/|__|       
     \/                   
   _____  .__  .__        
  /  _  \ |  | |  |       
 /  /_\  \|  | |  |       
/    |    \  |_|  |__     
\____|__  /____/____/     
        \/                
___________               
\__    ___/___            
  |    | /  _ \           
  |    |(  <_> )          
  |____| \____/           
                          
  _________               
 /   _____/ ____   ____   
 \_____  \_/ __ \_/ __ \  
 /        \  ___/\  ___/  
/_______  /\___  >\___  > 
        \/     \/     \/  
'''

print(art)



def parse_args():
    parser = argparse.ArgumentParser(description="Translate files in a directory and its subdirectories.")
    parser.add_argument("input_dir", help="The path to the input directory.")
    parser.add_argument("output_dir", help="The path to the output directory.")
    parser.add_argument("--src_language", default="en", help="The source language (default: 'en').")
    parser.add_argument("--tgt_language", default="es", help="The target language (default: 'es').")
    return parser.parse_args()

# Initialize the model and tokenizer
args = parse_args()
model_name = f"Helsinki-NLP/opus-mt-{args.src_language}-{args.tgt_language}"
tokenizer = MarianTokenizer.from_pretrained(model_name)
model = MarianMTModel.from_pretrained(model_name)


#gpu
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)



# Translation function
def translate_text(text):
    # Convert non-string input to string
    if not isinstance(text, str):
        text = str(text)

    # If the text is empty or contains only whitespace, return it as-is
    if not text.strip():
        return text

    # Split the input text into chunks of approximately 100 words
    text_chunks = text.split(' ')
    chunk_size = 100
    chunks = [' '.join(text_chunks[i:i + chunk_size]) for i in range(0, len(text_chunks), chunk_size)]

    translated_chunks = []
    for chunk in chunks:
        print(f"Translating: {chunk[:50]}...")  # progress verbosity

        # Tokenize and translate the chunk
        tokenized_chunk = tokenizer(chunk, return_tensors="pt")
        input_ids = tokenized_chunk["input_ids"].to(device)  # Move input tensor to the device

        translated_tokens = model.generate(input_ids, max_length=720)
        translated_chunk = tokenizer.decode(translated_tokens[0], skip_special_tokens=True)

        print(f"Translated: {translated_chunk[:50]}...")  # further verbosity
        translated_chunks.append(translated_chunk)

    # Combine the translated chunks
    translated_text = ' '.join(translated_chunks)
    return translated_text



  
# File processing function
def process_files(input_dir, output_dir):
    files_to_process = []
    for root, _, files in os.walk(input_dir):
        for file in files:
            file_path = os.path.join(root, file)
            files_to_process.append(file_path)

    progress_bar = tqdm(total=len(files_to_process), desc="Processing files", unit="file")

    def process_file_with_progress(file_path):
        try:
            translate_file(file_path)
        except Exception as e:
            progress_bar.write(f"Error processing {file_path}: {e}")
        finally:
            progress_bar.update()
# set threadpoolexecuter max_workers=
    with ThreadPoolExecutor(max_workers=10) as executor:
        executor.map(process_file_with_progress, files_to_process)

    progress_bar.close()
    
    
    
    
# Get filey    
    

def translate_file(file_path):
    file_ext = os.path.splitext(file_path)[1]
    if file_ext == ".docx":
        translate_docx(file_path)
    elif file_ext == ".xlsx":
        translate_xlsx(file_path)
    elif file_ext == ".xls":
        translate_xls(file_path)
    elif file_ext == ".csv":
        process_csv_in_batches(file_path)
    elif file_ext == ".json":
        translate_json(file_path)
    elif file_ext == ".html":
        translate_html(file_path)
    elif file_ext == ".odt":
        translate_odt(file_path)
    elif file_ext == ".txt":
        translate_txt(file_path)
    elif file_ext == ".rtf":
        translate_rtf(file_path)       
    elif file_ext == ".doc":
        translate_doc(file_path) 
    elif file_ext == ".eml":
        translate_eml(file_path)        
    else:
        print(f"Unsupported file type: {file_ext}")

        
# file handlers

def translate_dict(obj):
    if isinstance(obj, dict):
        return {key: translate_dict(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [translate_dict(item) for item in obj]
    elif isinstance(obj, str):
        return translate_text(obj)
    else:
        return obj

     
     
     
     
     
     
     
   
# Define translation functions for each file type


#rtf handling
def translate_rtf(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            rtf_content = file.read()
    except UnicodeDecodeError:
        with open(file_path, "rb") as file:
            rawdata = file.read()
        encoding = chardet.detect(rawdata)['encoding']
        try:
            rtf_content = rawdata.decode(encoding)
        except Exception as e:
            print(f"Chardet failed with error {e}. Defaulting to ISO-8859-1 decoding.")
            rtf_content = rawdata.decode('ISO-8859-1')

    text = rtf_to_text(rtf_content)

    # Clean up the text by removing redundant whitespaces
    text = re.sub(' +', ' ', text)
    
    translated_text = translate_text(text)

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir) + '.txt')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as output_file:
        output_file.write(translated_text)





#end rtf handling

#doc handling        
        
def doc_to_docx(file_path):
    docx_file_path = os.path.splitext(file_path)[0] + '.docx'

    # Run LibreOffice in headless mode to convert the .doc file to .docx
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'docx', file_path, '--outdir', os.path.dirname(file_path)])

    return docx_file_path

def translate_doc(file_path):
    # Convert .doc to .docx
    docx_file_path = doc_to_docx(file_path)
    
    # Translate the .docx file
    translate_docx(docx_file_path)

    # Remove the .docx file
    os.remove(docx_file_path)       
#end doc handling        


def translate_eml(file_path):
    # Open the .eml file and parse it into a EmailMessage object
    with open(file_path, "r") as file:
        msg = email.message_from_file(file, policy=policy.default)

    # Translate the subject and body of the email
    if msg['Subject']:
        msg['Subject'] = translate_text(msg['Subject'])
    if msg.get_body():
        msg.set_content(translate_text(msg.get_body().get_content()))

    # Save the translated email to the output directory
    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w") as output_file:
        output_file.write(msg.as_string())






def translate_docx(file_path):
    doc = Document(file_path)

    # Translate all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            try:
                run.text = translate_text(run.text)
            except Exception as e:
                print(f"Error translating paragraph in {file_path}: {e}")

    # Translate all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                try:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = translate_text(run.text)
                except Exception as e:
                    print(f"Error translating table cell in {file_path}: {e}")

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)



def translate_xlsx(file_path):
    workbook = openpyxl.load_workbook(file_path)
    for sheet in workbook.worksheets:
        for row in sheet:
            for cell in row:
                if isinstance(cell.value, str):
                    translated_text = translate_text(cell.value)
                    cell.value = translated_text

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    workbook.save(output_path)

import pandas as pd

def process_csv_in_batches(file_path, batch_size=100):
    # Set up a loop to process the CSV file in smaller batches
    output_rows = []
    for chunk in pd.read_csv(file_path, chunksize=batch_size):
        # Translate the rows in the current chunk
        translated_chunk = chunk.applymap(translate_text)
        output_rows.append(translated_chunk)
    
    # Concatenate the translated chunks and save the output
    output_df = pd.concat(output_rows)
    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    output_df.to_csv(output_path, index=False)


def translate_json(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        data = json.load(file)

    translated_data = translate_dict(data)

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as file:
        json.dump(translated_data, file, ensure_ascii=False, indent=2)

def translate_html(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file.read(), "html.parser")

    translate_soup(soup)

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(str(soup))


def translate_xls(file_path):
    workbook = xlrd.open_workbook(file_path)
    translated_workbook = xlwt.Workbook()

    for sheet_idx in range(workbook.nsheets):
        sheet = workbook.sheet_by_index(sheet_idx)
        translated_sheet = translated_workbook.add_sheet(sheet.name)

        for row_idx in range(sheet.nrows):
            for col_idx in range(sheet.ncols):
                cell_value = sheet.cell_value(row_idx, col_idx)

                if isinstance(cell_value, str):
                    translated_text = translate_text(cell_value)
                    translated_sheet.write(row_idx, col_idx, translated_text)
                else:
                    translated_sheet.write(row_idx, col_idx, cell_value)

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    translated_workbook.save(output_path)

    
def translate_odt(file_path):
    document = OpenDocumentText(file_path)
    translated_document = OpenDocumentText()

    for element in document.getElementsByType(P):
        paragraph = P()
        for child in element.childNodes:
            if child.qname[1] == "span":
                translated_text = translate_text(teletype.extractText(child))
                span = Span()
                span.addText(translated_text)
                paragraph.addElement(span)
            elif isinstance(child, str):
                translated_text = translate_text(child)
                paragraph.addText(translated_text)
        translated_document.text.addElement(paragraph)

    output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    translated_document.save(output_path)   
    
    


def translate_txt(file_path):
    try:
        # Detect the file encoding
        rawdata = open(file_path, 'rb').read()
        result = chardet.detect(rawdata)
        file_encoding = result['encoding']

        # Open the file with the detected encoding
        with open(file_path, "r", encoding=file_encoding) as file:
            content = file.read()

        translated_content = translate_text(content)

        output_path = os.path.join(output_dir, os.path.relpath(file_path, input_dir))
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "w", encoding='utf-8') as output_file:  # saving the translated text as utf-8
            output_file.write(translated_content)
    except Exception as e:
        print(f"Error processing {file_path}: {e}")

        
        
        
        
        
# Rest of the code remains the same

if __name__ == "__main__":
    input_dir = args.input_dir
    output_dir = args.output_dir

    if not os.path.exists(input_dir):
        print(f"Input directory '{input_dir}' does not exist.")
        sys.exit(1)

    os.makedirs(output_dir, exist_ok=True)
    process_files(input_dir, output_dir)
